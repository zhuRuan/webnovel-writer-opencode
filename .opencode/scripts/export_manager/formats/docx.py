#!/usr/bin/env python3
"""DOCX export — builds Word documents from chapter AST via python-docx."""
from __future__ import annotations

from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from export_manager.parser import md_to_blocks, ast_to_text


def export_docx(
    chapters: list,
    output_path: Path,
    title: str = "",
    author: Optional[str] = None,
) -> None:
    """Export chapters as a .docx file with Chinese novel typesetting."""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

    style = doc.styles["Normal"]
    style.font.size = Pt(12)
    style.font.name = "宋体"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    pf = style.paragraph_format
    pf.first_line_indent = Pt(24)
    pf.line_spacing = 1.8
    pf.space_after = Pt(4)

    if title:
        h = doc.add_heading(title, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if author:
        doc.add_paragraph(f"作者: {author}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    for ch in chapters:
        if hasattr(ch, 'path'):
            text = ch.path.read_text(encoding="utf-8")
            ch_index = ch.index
            ch_title = ch.title
        else:
            # tuple compat: (index, title, path)
            ch_index, ch_title, ch_path = ch[0], ch[1], ch[2]
            text = ch_path.read_text(encoding="utf-8")

        doc.add_page_break()
        heading = doc.add_heading(f"第{ch_index}章 {ch_title}", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        try:
            blocks = md_to_blocks(text)
        except Exception:
            doc.add_paragraph(text)
            continue

        for block in blocks:
            if not isinstance(block, dict):
                continue
            btype = block.get("type", "")

            if btype == "heading":
                continue
            elif btype == "thematic_break":
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run("* * *")
            elif btype == "paragraph":
                children = block.get("children", [])
                if not children:
                    continue
                p = doc.add_paragraph()
                _add_inline_runs(p, children)
            elif btype in ("blank_line", "block_code"):
                text_content = ast_to_text(block)
                if text_content.strip():
                    doc.add_paragraph(text_content)
            else:
                text_content = ast_to_text(block)
                if text_content.strip():
                    doc.add_paragraph(text_content)

    doc.save(str(output_path))


def _add_inline_runs(paragraph, children: list) -> None:
    """Add runs to paragraph, handling strong/bold inline formatting."""
    for child in children:
        if not isinstance(child, dict):
            continue
        ctype = child.get("type", "")
        if ctype == "text":
            paragraph.add_run(child.get("raw", ""))
        elif ctype == "strong":
            for inner in child.get("children", []):
                if isinstance(inner, dict) and inner.get("type") == "text":
                    run = paragraph.add_run(inner.get("raw", ""))
                    run.bold = True
        elif ctype == "emphasis":
            for inner in child.get("children", []):
                if isinstance(inner, dict) and inner.get("type") == "text":
                    run = paragraph.add_run(inner.get("raw", ""))
                    run.italic = True
        else:
            text = ast_to_text(child)
            if text:
                paragraph.add_run(text)
